"""
extractor.py — Extract raw text from PDF and DOCX files.
"""

import io


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """
    Extract text from a PDF using pdfplumber with pdfminer fallback.
    """
    text = ""
    try:
        import pdfplumber
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        if text.strip():
            return text.strip()
    except Exception as e:
        print(f"[pdfplumber] Error: {e}")

    # Fallback: pypdf
    try:
        import pypdf
        reader = pypdf.PdfReader(io.BytesIO(file_bytes))
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
        return text.strip()
    except Exception as e:
        print(f"[pypdf] Error: {e}")

    return text.strip()


def extract_text_from_docx(file_bytes: bytes) -> str:
    """
    Extract text from a Word (.docx) document.
    """
    try:
        import docx
        doc = docx.Document(io.BytesIO(file_bytes))
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)
        return "\n".join(paragraphs).strip()
    except Exception as e:
        print(f"[docx] Error: {e}")
        return ""
